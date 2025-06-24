from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from md2docx.utils import color_to_rgb

class DocStyle(object):
    def __init__(
            self,
            document,
            default_font='Times New Roman',
            chinease_font='宋体',
            font_size=12,
            heading_color='17a2b8'
        ):
        self.default_font = default_font
        self.chinease_font = chinease_font
        self.font_size = font_size
        self.heading_color = heading_color
        self.document = document

    def set_document_normal_style(self):
        style = self.document.styles['Normal']
        style.font.name = self.default_font
        style.font.size = Pt(self.font_size)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.chinease_font)

    def set_document_heading_style(self, bold=True, italic=False): 
        for paragraph in self.document.paragraphs:

            # 删除段落间距
            # paragraph.paragraph_format.space_before = Pt(0)
            # paragraph.paragraph_format.space_after = Pt(0)

            # 设置1.5倍行距
            paragraph.paragraph_format.line_spacing = 1.5

            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name.split()[-1])
                font_size = self.font_size - 2 + (6 - level) * 2
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(font_size)
                    font.color.rgb = color_to_rgb(self.heading_color)
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:ascii'), self.default_font)
                    rFonts.set(qn('w:hAnsi'), self.default_font)
                    rFonts.set(qn('w:eastAsia'), self.chinease_font)
                    font.bold = bold
                    font.italic = italic

    def set_image_style(self, max_width_inch=6):
        """Set image width to max_width_inch if it's larger than max_width_inch
        """
        for shape in self.document.inline_shapes:
            origin_width = shape.width
            origin_height = shape.height
            if origin_width > Inches(max_width_inch):
                ratio = Inches(max_width_inch) / origin_width
                shape.width = Inches(max_width_inch)
                shape.height = int(origin_height * ratio)

    def set_table_style(self):
        for table in self.document.tables:
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    self._set_cell_border(cell, color='FFFFFF')  # 加边框
                    if row_idx % 2 == 1:   # 交替背景色
                        self._set_cell_background(cell, "F2F2F2")  # 浅灰
                    else:
                        self._set_cell_background(cell, "c7e2f0")  # 浅蓝 

                    self._set_cell_paragraph(cell)

                    # self._set_cell_align(cell, vertical='center') # 垂直居中无效
    @classmethod
    def _set_cell_paragraph(cls, cell):
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            paragraph.paragraph_format.line_spacing = Pt(14)

    @classmethod
    def _set_cell_border(cls, cell, size="4", color="000000"):
        """
        设置单元格边框，直接操作底层XML
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        borders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), size)
            edge_el.set(qn('w:color'), color)
            edge_el.set(qn('w:space'), "0")
            borders.append(edge_el)

        tcPr.append(borders)

    @classmethod
    def _set_cell_background(cls, cell, fill_color: str):
        """
        设置单元格背景色，fill_color 为 6位RGB十六进制，如 F2F2F2
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)

        tcPr.append(shd)

    @classmethod
    def _set_cell_align(cls, cell, vertical='center'):
        """
        设置单元格对齐方式，vertical 为 'top', 'center', 'bottom'
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # 删除旧的 vAlign 元素
        for el in tcPr.findall(qn('w:vAlign')):
            tcPr.remove(el)

        # 创建新的 vAlign 元素
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), vertical)
        tcPr.append(vAlign)

    @classmethod
    def _set_run_background(cls, run, fill_color: str):
        """
        设置文本背景色，fill_color 为 6位RGB十六进制，如 F2F2F2
        """
        rPr = run._element.get_or_add_rPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        rPr.append(shd)

    @classmethod
    def _set_paragraph_background(cls, paragraph, hex_color='F2F2F2'):
        """设置整个段落的背景色"""
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')        # 填充类型
        shd.set(qn('w:color'), 'auto')       # 文字颜色
        shd.set(qn('w:fill'), hex_color)     # 背景色（十六进制）
        p_pr.append(shd)

    
    def set_styles(self):
        self.set_document_normal_style()
        self.set_document_heading_style()
        self.set_image_style()
        self.set_table_style()



