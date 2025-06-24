import os
import re
import tempfile
from docx import Document
from PIL import Image

from docx.shared import Inches
from typing import Optional, Tuple

from md2docx import utils


def process_image(
        doc: Document,
        line: str,
    ):
    """处理图像"""
    # 提取图像信息
    match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
    if not match:
        return
    
    alt_text = match.group(1)
    image_path = match.group(2)
    
    image = utils.ImageLoader.load_image(image_path)
    p = doc.add_paragraph()
    if image is None:
        run = p.add_run(f"[图像: {alt_text}]")
        run.font.italic = True
    else:
        width, height = utils.calculate_image_size(image)
        run = p.add_run()

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
            image.save(temp_file.name, format='PNG')
            run.add_picture(temp_file.name, width=width, height=height)

        try:
            os.unlink(temp_file.name)
        except:
            pass