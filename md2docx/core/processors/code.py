import os

from docx import Document
from docx.shared import Pt

from md2docx import utils



def add_code_block(doc: Document, code_text: str, background_color: str = 'F2F2F2') -> None:
    p = doc.add_paragraph()

    utils.DocStyle._set_paragraph_background(p, background_color)  # 设置段落背景色

    # 设置段落格式
    # p.paragraph_format.left_indent = Pt(20)
    # p.paragraph_format.right_indent = Pt(20)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 添加代码文本
    run = p.add_run(code_text)
    run.font.size = Pt(10)


def add_mermaid_code_block(doc: Document, code_text: str) -> None:
    image_path = utils.mermaid_to_png(code_text)
    if image_path:
        width, height = utils.calculate_image_size(image_path)
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(image_path, width=width, height=height)
        os.unlink(image_path)
        return True


def process_code_block(
        doc: Document,
        lines: list,
        start_idx: int,
        render_mermaid: bool = True,
    ) -> int:
    """处理代码块"""
    # 提取语言信息
    first_line = lines[start_idx].strip()
    language = first_line[3:].strip() if len(first_line) > 3 else ''

    # 收集代码内容
    code_lines = []
    i = start_idx + 1
    while i < len(lines) and not lines[i].strip().startswith('```'):
        code_lines.append(lines[i])
        i += 1
    
    # 创建代码块段落
    code_text = '\n'.join(code_lines)

    if language.lower() == 'mermaid' and  render_mermaid:
        if not add_mermaid_code_block(doc, code_text):
            add_code_block(doc, code_text)
    else:
        add_code_block(doc, code_text)
    
    return i + 1  # 跳过结束的```


def process_inline_code(doc: Document, text: str) -> str:
    """预处理行内代码，为后续处理做标记"""
    return text
