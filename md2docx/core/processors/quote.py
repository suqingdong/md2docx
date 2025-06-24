import re
from docx import Document
from docx.shared import RGBColor, Pt

from md2docx import utils

from .paragraph import add_hyperlink


def process_quote(
        doc: Document,
        line: str,
        left_indent: int = 24,
        italic: bool=True,
        color: str = 'FF0000'  # 默认灰色引用颜色
    ):
        """处理引用"""
        line = line.strip()[1:].strip()  # 移除 >
        # p = doc.add_paragraph(text)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(left_indent)


        # 提取 Markdown 超链接并替换为占位符
        link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
        link_matches = list(link_pattern.finditer(line))
        link_placeholders = {}

        for i, match in enumerate(link_matches):
            text, url = match.group(1), match.group(2)
            placeholder = f'[[[LINK_{i}]]]'
            link_placeholders[placeholder] = (text, url)
            line = line.replace(match.group(0), placeholder, 1)


        parts = re.split(r'(\[\[\[LINK_\d+\]\]\])', line)
        for part in parts:
            if not part:
                continue
            elif re.match(r'\[\[\[LINK_\d+\]\]\]', part):
                text, url = link_placeholders[part]
                add_hyperlink(p, text, url)
            else:
                p.add_run(part)
        
        # 设置引用样式
        for run in p.runs:
            run.font.italic = italic
            # run.font.color.rgb = utils.color_to_rgb(color)
            # run.font.color.rgb = RGBColor.from_string(color)
            run.font.color.rgb = RGBColor(150, 150, 150)
