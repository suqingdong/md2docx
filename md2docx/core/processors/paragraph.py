import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from md2docx.utils import DocStyle, color_to_rgb


def add_hyperlink(paragraph, text, url):
    """在 Word 段落中插入超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    # 蓝色 + 下划线
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    r_pr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    r_pr.append(u)

    new_run.append(r_pr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# 提取出的新函数：处理所有内联样式
def process_inline_styles(
        p: Paragraph,  # 注意：这里接收一个 Paragraph 对象，而不是 Document
        line: str,
        inline_color: str = 'C72E96',
        inline_background: str = 'F2F2F2',
    ):
    """处理段落内的加粗、斜体、行内代码、超链接"""

    # 移除HTML标签
    line = re.sub(r'<[^>]+>', '', line)

    # 特殊链接处理
    # [[xxx]](https://www.baidu.com) ==> [xxx]
    line = re.sub(r'\[\[([^\]]+)\]\]\(([^)]+)\)', r'[\1]', line)

    
    # 提取 Markdown 超链接并替换为占位符
    link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    link_matches = list(link_pattern.finditer(line))
    link_placeholders = {}

    for i, match in enumerate(link_matches):
        text, url = match.group(1), match.group(2)
        placeholder = f'[[[LINK_{i}]]]'
        link_placeholders[placeholder] = (text, url)
        line = line.replace(match.group(0), placeholder, 1)

    # 再处理加粗、斜体、代码等
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[\[\[LINK_\d+\]\]\])', line)

    for part in parts:
        if not part:
            continue
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = color_to_rgb(inline_color)
            DocStyle._set_run_background(run, inline_background)
        elif re.match(r'\[\[\[LINK_\d+\]\]\]', part):
            text, url = link_placeholders[part]
            add_hyperlink(p, text, url)
        else:
            p.add_run(part)


# 修改后的 process_paragraph
def process_paragraph(
        doc: Document,
        line: str,
        inline_color: str = 'C72E96',
        inline_background: str = 'F2F2F2',
    ):
    """处理普通段落，支持加粗、斜体、行内代码、超链接"""
    p = doc.add_paragraph()  # 1. 创建段落
    
    # 2. 调用新函数来填充内容
    process_inline_styles(p, line, inline_color, inline_background)
